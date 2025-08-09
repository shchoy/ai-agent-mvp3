"""
File handling utilities for processing ZIP files and documents
"""

import os
import zipfile
import tempfile
import shutil
from typing import List, Dict, Tuple, Optional
from pathlib import Path
import logging

# Document processing imports
import PyPDF2
from docx import Document
import pandas as pd
from openpyxl import load_workbook

from utils.config import Config

logger = logging.getLogger(__name__)


class FileHandler:
    """Handles file operations for tender documents"""
    
    def __init__(self):
        self.config = Config()
        self.temp_dirs = []
    
    def extract_zip_file(self, zip_path: str) -> str:
        """
        Extract ZIP file to temporary directory
        
        Args:
            zip_path: Path to the ZIP file
            
        Returns:
            Path to the temporary extraction directory
        """
        try:
            # Create temporary directory
            temp_dir = tempfile.mkdtemp()
            self.temp_dirs.append(temp_dir)
            
            # Extract ZIP file
            with zipfile.ZipFile(zip_path, 'r') as zip_ref:
                zip_ref.extractall(temp_dir)
            
            logger.info(f"Extracted ZIP file to: {temp_dir}")
            return temp_dir
            
        except zipfile.BadZipFile:
            raise ValueError("Invalid ZIP file provided")
        except Exception as e:
            raise Exception(f"Error extracting ZIP file: {str(e)}")
    
    def find_document_files(self, directory: str) -> List[str]:
        """
        Find all supported document files in the directory
        
        Args:
            directory: Directory to search
            
        Returns:
            List of file paths
        """
        document_files = []
        
        for root, dirs, files in os.walk(directory):
            for file in files:
                # Skip macOS system files and hidden files
                if self._should_skip_file(file):
                    logger.debug(f"Skipping system/hidden file: {file}")
                    continue
                    
                file_path = os.path.join(root, file)
                file_ext = Path(file).suffix.lower()
                
                if file_ext in self.config.SUPPORTED_EXTENSIONS:
                    # Check file size
                    if os.path.getsize(file_path) <= self.config.MAX_FILE_SIZE:
                        document_files.append(file_path)
                    else:
                        logger.warning(f"File too large, skipping: {file_path}")
        
        logger.info(f"Found {len(document_files)} supported documents")
        return document_files
    
    def _should_skip_file(self, filename: str) -> bool:
        """
        Check if file should be skipped (system files, hidden files, etc.)
        
        Args:
            filename: Name of the file
            
        Returns:
            True if file should be skipped
        """
        skip_patterns = [
            # macOS system files
            '.__MACOSX',
            '._',
            '.DS_Store',
            # Hidden files
            '.',
            # Thumbs.db (Windows)
            'Thumbs.db',
            'desktop.ini'
        ]
        
        for pattern in skip_patterns:
            if filename.startswith(pattern):
                return True
        
        return False
    
    def extract_text_from_pdf(self, file_path: str) -> str:
        """Extract text from PDF file with robust error handling"""
        try:
            # Skip if it's a system file
            filename = os.path.basename(file_path)
            if self._should_skip_file(filename):
                logger.debug(f"Skipping system file: {filename}")
                return ""
            
            text = ""
            with open(file_path, 'rb') as file:
                try:
                    pdf_reader = PyPDF2.PdfReader(file)
                    
                    # Check if PDF is encrypted
                    if pdf_reader.is_encrypted:
                        logger.warning(f"PDF is encrypted, skipping: {file_path}")
                        return ""
                    
                    # Extract text from all pages
                    for page_num, page in enumerate(pdf_reader.pages):
                        try:
                            page_text = page.extract_text()
                            if page_text:
                                text += page_text + "\n"
                        except Exception as page_error:
                            logger.warning(f"Error extracting page {page_num} from {file_path}: {str(page_error)}")
                            continue
                    
                except PyPDF2.errors.PdfReadError as pdf_error:
                    logger.error(f"PDF read error for {file_path}: {str(pdf_error)}")
                    return ""
                except Exception as read_error:
                    logger.error(f"Unexpected error reading PDF {file_path}: {str(read_error)}")
                    return ""
            
            return text.strip()
            
        except FileNotFoundError:
            logger.error(f"PDF file not found: {file_path}")
            return ""
        except PermissionError:
            logger.error(f"Permission denied reading PDF: {file_path}")
            return ""
        except Exception as e:
            logger.error(f"Error reading PDF {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_docx(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        try:
            doc = Document(file_path)
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Also extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += cell.text + " "
                    text += "\n"
            
            return text.strip()
        except Exception as e:
            logger.error(f"Error reading DOCX {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_excel(self, file_path: str) -> str:
        """Extract text from Excel file"""
        try:
            # Try with pandas first
            try:
                df = pd.read_excel(file_path, sheet_name=None)
                text = ""
                for sheet_name, sheet_df in df.items():
                    text += f"Sheet: {sheet_name}\n"
                    text += sheet_df.to_string() + "\n\n"
                return text.strip()
            except:
                # Fallback to openpyxl
                workbook = load_workbook(file_path, data_only=True)
                text = ""
                for sheet_name in workbook.sheetnames:
                    sheet = workbook[sheet_name]
                    text += f"Sheet: {sheet_name}\n"
                    for row in sheet.iter_rows(values_only=True):
                        text += " ".join([str(cell) if cell is not None else "" for cell in row]) + "\n"
                    text += "\n"
                return text.strip()
        except Exception as e:
            logger.error(f"Error reading Excel {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_txt(self, file_path: str) -> str:
        """Extract text from TXT file"""
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                return file.read().strip()
        except UnicodeDecodeError:
            # Try with different encoding
            try:
                with open(file_path, 'r', encoding='latin-1') as file:
                    return file.read().strip()
            except Exception as e:
                logger.error(f"Error reading TXT {file_path}: {str(e)}")
                return ""
        except Exception as e:
            logger.error(f"Error reading TXT {file_path}: {str(e)}")
            return ""
    
    def extract_text_from_file(self, file_path: str) -> Tuple[str, str]:
        """
        Extract text from any supported file type
        
        Args:
            file_path: Path to the file
            
        Returns:
            Tuple of (filename, extracted_text)
        """
        filename = os.path.basename(file_path)
        
        # Skip system files
        if self._should_skip_file(filename):
            logger.debug(f"Skipping system file: {filename}")
            return filename, ""
        
        file_ext = Path(file_path).suffix.lower()
        
        logger.info(f"Extracting text from: {filename}")
        
        try:
            if file_ext == '.pdf':
                text = self.extract_text_from_pdf(file_path)
            elif file_ext in ['.docx', '.doc']:
                text = self.extract_text_from_docx(file_path)
            elif file_ext in ['.xlsx', '.xls']:
                text = self.extract_text_from_excel(file_path)
            elif file_ext == '.txt':
                text = self.extract_text_from_txt(file_path)
            else:
                logger.warning(f"Unsupported file type: {file_ext}")
                text = ""
            
            if not text or not text.strip():
                logger.warning(f"No text extracted from: {filename}")
                
            return filename, text
            
        except Exception as e:
            logger.error(f"Unexpected error processing file {filename}: {str(e)}")
            return filename, ""
    
    def process_zip_file(self, zip_path: str) -> Dict[str, str]:
        """
        Process entire ZIP file and extract text from all documents
        
        Args:
            zip_path: Path to the ZIP file
            
        Returns:
            Dictionary mapping filename to extracted text
        """
        try:
            # Extract ZIP file
            extraction_dir = self.extract_zip_file(zip_path)
            
            # Get all document files
            document_files = self.find_document_files(extraction_dir)
            
            if not document_files:
                logger.warning("No supported document files found in ZIP archive")
                return {}
            
            # Extract text from each file
            extracted_texts = {}
            successfully_processed = 0
            
            for file_path in document_files:
                filename, text = self.extract_text_from_file(file_path)
                if text and text.strip():  # Only include files with actual text content
                    extracted_texts[filename] = text
                    successfully_processed += 1
                else:
                    logger.warning(f"No usable text extracted from: {filename}")
            
            logger.info(f"Successfully processed {successfully_processed}/{len(document_files)} documents")
            
            if not extracted_texts:
                logger.error("No text could be extracted from any documents in the ZIP file")
            
            return extracted_texts
            
            logger.info(f"Successfully extracted text from {len(extracted_texts)} files")
            return extracted_texts
            
        except Exception as e:
            logger.error(f"Error processing ZIP file: {str(e)}")
            raise
    
    def cleanup(self):
        """Clean up temporary directories"""
        for temp_dir in self.temp_dirs:
            try:
                shutil.rmtree(temp_dir)
                logger.info(f"Cleaned up temporary directory: {temp_dir}")
            except Exception as e:
                logger.warning(f"Failed to clean up {temp_dir}: {str(e)}")
        self.temp_dirs.clear()
    
    def __del__(self):
        """Cleanup on destruction"""
        self.cleanup()
